from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfWriter

from src.ingest.loaders import load_directory, load_docx, load_html, load_pdf
from src.ingest.splitters import split_recursive


def test_load_pdf_preserves_source_and_page_metadata(tmp_path: Path):
    pdf_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with pdf_path.open("wb") as f:
        writer.write(f)

    docs = load_pdf(pdf_path)

    assert len(docs) == 1
    assert docs[0].metadata["source"] == str(pdf_path)
    assert docs[0].metadata["page"] == 0


def test_load_docx_returns_text_with_source_metadata(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_heading("KnowledgeOps", level=1)
    document.add_paragraph("Sprint 1 ingests Word documents.")
    document.save(docx_path)

    docs = load_docx(docx_path)

    assert len(docs) == 1
    assert "KnowledgeOps" in docs[0].page_content
    assert "Sprint 1 ingests Word documents." in docs[0].page_content
    assert docs[0].metadata["source"] == str(docx_path)


def test_load_html_returns_visible_text_with_source_metadata(tmp_path: Path):
    html_path = tmp_path / "sample.html"
    html_path.write_text(
        "<html><head><title>Ignored</title><script>hidden()</script></head>"
        "<body><main><h1>Research Loop</h1><p>HTML evidence is loaded.</p></main></body></html>",
        encoding="utf-8",
    )

    docs = load_html(html_path)

    assert len(docs) == 1
    assert "Research Loop" in docs[0].page_content
    assert "HTML evidence is loaded." in docs[0].page_content
    assert "hidden()" not in docs[0].page_content
    assert docs[0].metadata["source"] == str(html_path)


def test_load_directory_dispatches_supported_local_documents(tmp_path: Path):
    html_path = tmp_path / "sample.html"
    html_path.write_text("<body><p>Directory HTML evidence.</p></body>", encoding="utf-8")
    docx_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Directory Word evidence.")
    document.save(docx_path)
    (tmp_path / "ignored.txt").write_text("not supported", encoding="utf-8")

    docs = load_directory(tmp_path)

    sources = {doc.metadata["source"] for doc in docs}
    assert str(html_path) in sources
    assert str(docx_path) in sources
    assert all("ignored.txt" not in source for source in sources)


def test_split_recursive_uses_sprint1_overlap_alias_and_preserves_source_metadata():
    docs = [
        Document(
            page_content="alpha beta gamma delta epsilon zeta eta theta iota kappa",
            metadata={"source": "fixture.md"},
        )
    ]

    chunks = split_recursive(docs, chunk_size=30, overlap=5)

    assert len(chunks) >= 2
    assert all(len(chunk.page_content) <= 30 for chunk in chunks)
    assert all(chunk.metadata["source"] == "fixture.md" for chunk in chunks)
