import os
from pathlib import Path
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".html", ".htm"}
EXCLUDED_DIRECTORY_NAMES = {
    "__pycache__",
    "artifacts",
    "faiss",
    "langfuse-clickhouse",
    "langfuse-clickhouse-data",
    "langfuse-clickhouse-logs",
    "langfuse-minio",
    "langfuse-postgres",
    "langfuse-redis",
    "milvus",
}


def _normalize_source(path_or_url: str | Path) -> str:
    return str(path_or_url)


def _ensure_file(path: str | Path) -> Path:
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Local document does not exist: {file_path}")
    if not file_path.is_file():
        raise ValueError(f"Expected a file path, got: {file_path}")
    return file_path


def load_pdf(path: str | Path) -> list[Document]:
    """加载 PDF → List[Document]，每页一个 Document"""
    file_path = _ensure_file(path)
    docs = PyPDFLoader(str(file_path)).load()
    for page_no, doc in enumerate(docs):
        doc.metadata["source"] = str(file_path)
        doc.metadata.setdefault("page", page_no)
    return docs


def load_docx(path: str | Path) -> list[Document]:
    """加载 Word 文档 → List[Document]"""
    file_path = _ensure_file(path)
    document = DocxDocument(str(file_path))
    text = "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": str(file_path), "document_type": "docx"},
        )
    ]


def load_html(path_or_url: str | Path) -> list[Document]:
    """加载本地 HTML 文件或 URL → List[Document]"""
    source = _normalize_source(path_or_url)
    parsed = urlparse(source)
    if parsed.scheme in {"http", "https"}:
        response = httpx.get(source, timeout=10.0)
        response.raise_for_status()
        html = response.text
    else:
        file_path = _ensure_file(path_or_url)
        source = str(file_path)
        html = file_path.read_text(encoding="utf-8")

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    container = soup.find("main") or soup.find("article") or soup.body or soup
    text = container.get_text(separator="\n", strip=True)
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": source, "document_type": "html"},
        )
    ]


def load_url(url: str) -> list[Document]:
    """加载网页 → List[Document]"""
    return load_html(url)


def load_directory(directory: str | Path, glob: str = "**/*") -> list[Document]:
    """批量加载一个目录下所有支持的文档"""
    root = Path(directory)
    if not root.exists():
        raise FileNotFoundError(f"Local ingest directory does not exist: {root}")
    if not root.is_dir():
        raise ValueError(f"Expected a directory path, got: {root}")

    docs: list[Document] = []
    for file_path in _iter_supported_files(root, glob=glob):
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            docs.extend(load_pdf(file_path))
        elif suffix == ".docx":
            docs.extend(load_docx(file_path))
        elif suffix in {".html", ".htm"}:
            docs.extend(load_html(file_path))
    return docs


def _iter_supported_files(root: Path, glob: str) -> list[Path]:
    if glob == "**/*":
        candidates: list[Path] = []
        for current_root, dirnames, filenames in os.walk(root, topdown=True, onerror=lambda _: None):
            dirnames[:] = [
                dirname
                for dirname in dirnames
                if dirname not in EXCLUDED_DIRECTORY_NAMES
            ]
            current_path = Path(current_root)
            for filename in filenames:
                file_path = current_path / filename
                if file_path.suffix.lower() in SUPPORTED_SUFFIXES:
                    candidates.append(file_path)
        return sorted(candidates)

    candidates = []
    try:
        paths = sorted(root.glob(glob))
    except OSError:
        return []
    for file_path in paths:
        if _has_excluded_parent(file_path, root):
            continue
        if file_path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        try:
            if file_path.is_file():
                candidates.append(file_path)
        except OSError:
            continue
    return candidates


def _has_excluded_parent(path: Path, root: Path) -> bool:
    try:
        relative_parts = path.relative_to(root).parts
    except ValueError:
        relative_parts = path.parts
    return any(part in EXCLUDED_DIRECTORY_NAMES for part in relative_parts[:-1])
